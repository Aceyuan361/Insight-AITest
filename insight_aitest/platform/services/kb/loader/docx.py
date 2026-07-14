# -*- coding: utf-8 -*-
"""Word 文档 Loader（.docx）。用 python-docx 提取段落 + 表格文本。"""

from __future__ import annotations

from pathlib import Path

from insight_aitest.platform.services.kb.loader import register_loader
from insight_aitest.platform.services.kb.loader.base import DocumentLoader
from insight_aitest.platform.services.kb.models import ParsedDocument


@register_loader(".docx")
class DocxLoader(DocumentLoader):
    """Word .docx 文档 Loader。提取段落和表格文本。

    注意：全量加载到内存（python-docx 无流式模式），依赖上传端点的
    max_upload_mb 检查限制单文件大小（默认 20MB）。
    """

    def load(self, path: Path) -> ParsedDocument:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        parts: list[str] = []

        # 段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # 表格
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        content = "\n".join(parts)
        return ParsedDocument(filename=path.name, content=content, meta={"format": "docx"})
